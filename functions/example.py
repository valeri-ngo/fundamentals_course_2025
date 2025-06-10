from docx import Document
from fpdf import FPDF
import os
from zipfile import ZipFile

# Input data
full_name = "Valeri Lien Ngo"
phone = "+359 877 150 035"
email = "sacredlie@abv.bg"
location = "Ruse, Bulgaria (originally from Gabrovo, 20 years in Sofia)"
summary = (
    "Reliable and proactive logistics professional with over 10 years of experience in warehouse operations, "
    "driving, and dispatch coordination. Since 2019, serving as a driver/dispatcher with in-depth knowledge of "
    "company logistics and client locations. Strong team player with excellent computer skills, currently expanding "
    "expertise in programming (Python) at SoftUni. Actively seeking new challenges for both professional and personal development."
)
experience = [
    {
        "title": "Driver / Dispatcher / Warehouse Operator",
        "company": "SFK-Truck Ltd.",
        "location": "Sofia / Ruse, Bulgaria",
        "period": "2014 – Present (with 9-month break)",
        "details": [
            "Transitioned from warehouse operator to driver and dispatcher around 2019.",
            "Responsible for preparing goods for shipment, route planning, and client delivery.",
            "Coordinated the company’s vehicle fleet and assisted colleagues with logistics planning.",
            "Maintained stock accuracy and ensured timely deliveries.",
            "Gained extensive knowledge of client locations and logistics procedures.",
        ],
    },
    {
        "title": "Warehouse Worker",
        "company": "Inter Cars Bulgaria",
        "location": "Sofia, Bulgaria",
        "period": "Approx. 9 months (year unspecified)",
        "details": [
            "Performed loading/unloading operations in a high-paced warehouse environment.",
            "Managed stock organization and supported dispatch team.",
            "Complied with safety standards and helped optimize warehouse workflow.",
        ],
    },
]
education = "Evening Secondary School “Penyo Penev” – Sofia, Bulgaria\nCompleted Secondary Education (No Specific Major)"
courses = "Programming with Python – SoftUni (currently enrolled)\n- Learning programming fundamentals, software logic, and basic coding in Python."
skills = [
    "Driving (active driver since 2010)",
    "Dispatch coordination",
    "Warehouse logistics & stock handling",
    "Route planning & customer delivery",
    "Computer literate (MS Office, basic software skills)",
    "Python programming (beginner level)",
    "Excellent organizational skills",
    "Highly motivated and self-driven",
    "Strong teamwork and communication",
]
languages = "Bulgarian – Native\nEnglish – Intermediate"

# File paths
base_path = "/mnt/data/valeri_cv"
os.makedirs(base_path, exist_ok=True)
docx_path = os.path.join(base_path, "Valeri_Lien_Ngo_CV.docx")
pdf_path = os.path.join(base_path, "Valeri_Lien_Ngo_CV.pdf")
html_path = os.path.join(base_path, "Valeri_Lien_Ngo_CV.html")
zip_path = os.path.join("/mnt/data", "Valeri_Lien_Ngo_CV_Package.zip")

# Create Word Document
doc = Document()
doc.add_heading("Curriculum Vitae", level=1)
doc.add_paragraph(full_name)
doc.add_paragraph(f"Phone: {phone}\nEmail: {email}\nLocation: {location}")

doc.add_heading("Professional Summary", level=2)
doc.add_paragraph(summary)

doc.add_heading("Work Experience", level=2)
for job in experience:
    doc.add_paragraph(f"{job['title']}", style="Heading 3")
    doc.add_paragraph(f"{job['company']} – {job['location']}\n{job['period']}")
    for detail in job["details"]:
        doc.add_paragraph(f"- {detail}")

doc.add_heading("Education", level=2)
doc.add_paragraph(education)

doc.add_heading("Courses and Training", level=2)
doc.add_paragraph(courses)

doc.add_heading("Skills", level=2)
for skill in skills:
    doc.add_paragraph(f"- {skill}")

doc.add_heading("Languages", level=2)
doc.add_paragraph(languages)

doc.save(docx_path)

# Create PDF
pdf = FPDF()
pdf.add_page()
pdf.set_font("Arial", size=12)
pdf.cell(200, 10, txt="Curriculum Vitae", ln=True, align='C')
pdf.ln(10)
pdf.multi_cell(0, 10, f"{full_name}\nPhone: {phone}\nEmail: {email}\nLocation: {location}")
pdf.ln(5)
pdf.set_font("Arial", 'B', 12)
pdf.cell(0, 10, "Professional Summary", ln=True)
pdf.set_font("Arial", size=12)
pdf.multi_cell(0, 10, summary)
pdf.ln(5)
pdf.set_font("Arial", 'B', 12)
pdf.cell(0, 10, "Work Experience", ln=True)
pdf.set_font("Arial", size=12)
for job in experience:
    pdf.set_font("Arial", 'B', 12)
    pdf.cell(0, 10, job["title"], ln=True)
    pdf.set_font("Arial", size=12)
    pdf.multi_cell(0, 10, f"{job['company']} – {job['location']}\n{job['period']}")
    for detail in job["details"]:
        pdf.multi_cell(0, 10, f"- {detail}")
pdf.ln(5)
pdf.set_font("Arial", 'B', 12)
pdf.cell(0, 10, "Education", ln=True)
pdf.set_font("Arial", size=12)
pdf.multi_cell(0, 10, education)
pdf.ln(2)
pdf.set_font("Arial", 'B', 12)
pdf.cell(0, 10, "Courses and Training", ln=True)
pdf.set_font("Arial", size=12)
pdf.multi_cell(0, 10, courses)
pdf.ln(2)
pdf.set_font("Arial", 'B', 12)
pdf.cell(0, 10, "Skills", ln=True)
pdf.set_font("Arial", size=12)
for skill in skills:
    pdf.multi_cell(0, 10, f"- {skill}")
pdf.ln(2)
pdf.set_font("Arial", 'B', 12)
pdf.cell(0, 10, "Languages", ln=True)
pdf.set_font("Arial", size=12)
pdf.multi_cell(0, 10, languages)
pdf.output(pdf_path)

# Create HTML
html_content = f"""
<html>
<head><title>CV - Valeri Lien Ngo</title></head>
<body>
<h1>Curriculum Vitae</h1>
<p><strong>{full_name}</strong><br>
Phone: {phone}<br>
Email: {email}<br>
Location: {location}</p>

<h2>Professional Summary</h2>
<p>{summary}</p>

<h2>Work Experience</h2>
"""
for job in experience:
    html_content += f"<h3>{job['title']}</h3><p>{job['company']} – {job['location']}<br>{job['period']}</p><ul>"
    for detail in job["details"]:
        html_content += f"<li>{detail}</li>"
    html_content += "</ul>"

html_content += f"""
<h2>Education</h2>
<p>{education}</p>

<h2>Courses and Training</h2>
<p>{courses}</p>

<h2>Skills</h2>
<ul>
"""
for skill in skills:
    html_content += f"<li>{skill}</li>"
html_content += "</ul>"

html_content += f"""
<h2>Languages</h2>
<p>{languages}</p>
</body>
</html>
"""

with open(html_path, "w", encoding="utf-8") as f:
    f.write(html_content)

# Zip all files
with ZipFile(zip_path, 'w') as zipf:
    zipf.write(docx_path, arcname="Valeri_Lien_Ngo_CV.docx")
    zipf.write(pdf_path, arcname="Valeri_Lien_Ngo_CV.pdf")
    zipf.write(html_path, arcname="Valeri_Lien_Ngo_CV.html")

zip_path
